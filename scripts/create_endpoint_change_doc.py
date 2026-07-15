from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Permintaan_Informasi_Perubahan_Endpoint_API.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyper_text(paragraph, text, bold=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Permintaan Informasi Perubahan Endpoint API")
run.font.name = "Calibri"
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
subtitle_run = subtitle.add_run("Dokumen permintaan mapping endpoint baru atas perubahan URL integrasi API")
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor.from_string("555555")

doc.add_heading("Latar Belakang", level=1)
p = doc.add_paragraph()
p.add_run(
    "Sehubungan dengan adanya informasi perubahan URL API dari endpoint lama ke endpoint baru, "
    "kami membutuhkan konfirmasi endpoint baru yang perlu digunakan untuk masing-masing layanan. "
    "Dokumen ini disusun sebagai daftar pertanyaan dan referensi mapping agar proses penyesuaian "
    "integrasi dapat dilakukan dengan tepat."
)

doc.add_heading("Contoh Endpoint yang Sudah Diketahui", level=1)
p = doc.add_paragraph()
p.add_run("Contoh endpoint lama untuk register pemungut:")
p = doc.add_paragraph()
p.style = styles["Normal"]
r = p.add_run("http://36.94.117.75:8443/api/register/pengguna/add")
r.font.name = "Courier New"
r.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run("Contoh endpoint baru yang diterima sebagai referensi:")
p = doc.add_paragraph()
r = p.add_run("https://sim-integrasi.pajak.go.id/dgt/partners/mtd/v1.0/pemunggut")
r.font.name = "Courier New"
r.font.size = Pt(10)

doc.add_heading("Permintaan Konfirmasi Mapping Endpoint", level=1)
p = doc.add_paragraph()
p.add_run(
    "Mohon diinformasikan endpoint baru yang menggantikan setiap endpoint lama pada tabel berikut."
)

rows = [
    ("1", "http://36.94.117.75:8443/api/register/printer", ""),
    ("2", "http://36.94.117.75:8443/api/register/cartridge", ""),
    ("3", "http://36.94.117.75:8443/api/pengguna/validasi-pengguna?npwpPengguna=913448826065000", ""),
    ("4", "http://36.94.117.75:8443/api/register/pengguna/add", "https://sim-integrasi.pajak.go.id/dgt/partners/mtd/v1.0/pemunggut"),
    ("5", "http://36.94.117.75:8443/api/mon/pengguna/status?kodeDistributor=009&npwpPengguna=913448826065000", ""),
    ("6", "http://36.94.117.75:8443/api/pengguna/update_status?kodeDistributor=009&npwpPengguna=913448826065000", ""),
    ("7", "http://36.94.117.75:8443/api/register/pengguna/printer/add", ""),
    ("8", "http://36.94.117.75:8443/api/register/cartridge/update", ""),
    ("9", "http://36.94.117.75:8443/api/request/sn?token=E6DH9812GQ91ETSQLHKRRAIU8O78JL3TR2J6C1LB4N4DLBO2C6CG====", ""),
    ("10", "http://36.94.117.75:8443/api/pembubuhan/sn?token=NEC2TUF1OSB0B54CH0M0061NMPQ1S0DQ51PMT3F1NKQLTITDDJ0G====", ""),
    ("11", "http://36.94.117.75:8443/api/pembubuhan/update/sn?token=AUK8L212RP71DGS2537F28I5835BGSS8JHCHTBMQE68L6OIL3B4G====", ""),
    ("12", "http://36.94.117.75:8443/authenticate/token", ""),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, text in enumerate(("No", "Endpoint Lama", "Endpoint Baru")):
    hdr[i].text = text
    set_cell_shading(hdr[i], "F2F4F7")
    for paragraph in hdr[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
set_repeat_table_header(table.rows[0])

for no, old, new in rows:
    cells = table.add_row().cells
    cells[0].text = no
    cells[1].text = old
    cells[2].text = new
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

set_table_geometry(table, [720, 4540, 4100])

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if cell is not row.cells[0]:
                    run.font.name = "Courier New"

doc.add_paragraph()

doc.add_heading("Informasi Tambahan yang Dibutuhkan", level=1)
items = [
    "Apakah seluruh endpoint lama tersebut sudah memiliki padanan endpoint baru pada domain https://sim-integrasi.pajak.go.id/dgt/partners/mtd/v1.0?",
    "Apakah terdapat perubahan HTTP method, misalnya GET, POST, PUT, atau PATCH?",
    "Apakah terdapat perubahan format request body?",
    "Apakah terdapat perubahan parameter query?",
    "Apakah terdapat perubahan header atau skema autentikasi?",
    "Apakah terdapat perubahan format response?",
    "Apakah token atau credential lama masih dapat digunakan, atau perlu credential baru?",
    "Apakah terdapat perbedaan endpoint untuk environment development, staging, dan production?",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_heading("Catatan", level=1)
p = doc.add_paragraph()
p.add_run(
    "Pada contoh endpoint baru tertulis path "
)
r = p.add_run("/pemunggut")
r.font.name = "Courier New"
r.font.size = Pt(10)
p.add_run(
    ". Mohon konfirmasi apakah penulisan tersebut memang sudah sesuai dengan dokumentasi resmi, "
    "karena perbedaan penulisan path akan memengaruhi keberhasilan akses endpoint."
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("Permintaan Informasi Perubahan Endpoint API")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor.from_string("555555")

doc.save(OUT)
print(OUT)
